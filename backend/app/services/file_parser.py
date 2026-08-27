import subprocess
import logging
from pathlib import Path
from docx import Document

logger = logging.getLogger(__name__)


def _is_valid_text(text: str, threshold: float = 0.15) -> bool:
    """校验解码后文本是否为有效自然语言，控制字符占比超过阈值视为乱码"""
    if not text:
        return False
    control_chars = sum(1 for c in text if ord(c) < 32 and c not in '\n\r\t')
    return (control_chars / len(text)) < threshold


def _extract_table_with_merged_cells(table) -> str:
    """
    提取表格文本，处理纵向/横向合并单元格。
    返回以制表符分隔列、换行符分隔行的字符串。
    """
    rows = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            # 合并单元格内容会自然出现在第一个单元格，python-docx 默认返回空字符串给被合并的单元格
            row_cells.append(cell.text.strip())
        rows.append('\t'.join(row_cells))
    return '\n'.join(rows)


def _extract_header_footer(header_or_footer) -> list:
    """提取页眉或页脚中的所有段落文本"""
    texts = []
    for para in header_or_footer.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    return texts


def _extract_textboxes(doc) -> list:
    """提取文档中的所有文本框内容"""
    texts = []
    try:
        for elem in doc.element.body.iter():
            if elem.tag.endswith('txbxContent'):
                inner_text = ''.join(elem.itertext()).strip()
                if inner_text:
                    texts.append(inner_text)
    except Exception:
        pass
    return texts


def _extract_text_from_ole(file_path: str) -> str:
    """使用 olefile 提取 WordDocument 流中的 Unicode 文本，支持中文"""
    import olefile
    ole = olefile.OleFileIO(file_path)
    if not ole.exists('WordDocument'):
        return ""
    data = ole.openstream('WordDocument').read()
    # 尝试解析为 UTF-16LE（Word 内部存储 Unicode 常用方式）
    try:
        text = data.decode('utf-16-le', errors='ignore')
        clean = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
        if clean.strip() and len(clean.strip()) > 50:
            return clean
    except Exception:
        pass
    # 降级：提取 ASCII 及中文字符范围
    chars = []
    for i in range(0, len(data) - 1, 2):
        code = data[i] | (data[i + 1] << 8)
        if 0x4E00 <= code <= 0x9FFF or (32 <= code < 127):
            chars.append(chr(code))
    return ''.join(chars)


class DocumentParseError(Exception):
    """文档解析基类异常"""
    pass


class DocParseError(DocumentParseError):
    def __init__(self, file_path: str, attempted_parsers: list):
        self.file_path = file_path
        self.attempted_parsers = attempted_parsers
        msg = f"Unable to parse .doc file: {file_path}, attempted: {attempted_parsers}"
        super().__init__(msg)


class DocxParseError(DocumentParseError):
    def __init__(self, file_path: str, error: str):
        self.file_path = file_path
        self.error = error
        msg = f"Unable to parse .docx file: {file_path}, error: {error}"
        super().__init__(msg)


class TxtParseError(DocumentParseError):
    def __init__(self, file_path: str, attempted_encodings: list):
        self.file_path = file_path
        self.attempted_encodings = attempted_encodings
        msg = f"Unable to parse .txt file: {file_path}, attempted encodings: {attempted_encodings}"
        super().__init__(msg)


class PdfParseError(DocumentParseError):
    def __init__(self, file_path: str, error: str):
        self.file_path = file_path
        self.error = error
        msg = f"Unable to parse .pdf file: {file_path}, error: {error}"
        super().__init__(msg)


def read_pdf(file_path: Path) -> str:
    """读取 PDF 文件，提取所有页面的文本内容"""
    import fitz
    doc = fitz.open(file_path)
    texts = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            texts.append(text.strip())
    doc.close()
    content = '\n'.join(texts)
    if len(content.strip()) < 50:
        raise PdfParseError(str(file_path), "extracted content too short")
    return content

def read_docx(file_path: Path) -> str:
    """
    读取 .docx 文件，提取段落、表格（含合并处理）、页眉页脚、文本框。
    返回文本长度不足 50 字符时抛出异常。
    """
    try:
        doc = Document(file_path)
        full_text = []

        # 段落
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 表格（含合并单元格处理）
        for table in doc.tables:
            table_text = _extract_table_with_merged_cells(table)
            if table_text.strip():
                full_text.append(table_text)

        # 页眉页脚
        for section in doc.sections:
            full_text.extend(_extract_header_footer(section.header))
            full_text.extend(_extract_header_footer(section.footer))

        # 文本框
        full_text.extend(_extract_textboxes(doc))

        content = '\n'.join(full_text)
        if len(content.strip()) < 50:
            raise DocxParseError(str(file_path), "extracted content length < 50 chars")
        return content
    except DocxParseError:
        raise
    except Exception as e:
        raise DocxParseError(str(file_path), str(e))

def read_txt(file_path: Path) -> str:
    """
    读取 .txt 文件，尝试多种常见编码。
    解码后进行文本质量校验，全部失败则抛出 TxtParseError。
    """
    encodings = ['utf-8-sig', 'utf-8', 'gb18030', 'gbk', 'big5', 'utf-16', 'latin-1']
    attempted = []
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                content = f.read()
                if content and _is_valid_text(content):
                    if len(content.strip()) < 50:
                        attempted.append(f'{enc} (too short)')
                        continue
                    logger.debug(f"[read_txt] succeeded with {enc} for {file_path}")
                    return content
                attempted.append(f'{enc} (invalid content)')
        except (UnicodeDecodeError, UnicodeError) as e:
            attempted.append(f'{enc} ({type(e).__name__})')
            continue
        except Exception as e:
            attempted.append(f'{enc} ({str(e)})')
            continue

    logger.error(f"All encodings failed for {file_path}", extra={
        'file': str(file_path), 'attempted_encodings': attempted
    })
    raise TxtParseError(str(file_path), attempted)

def read_doc(file_path: Path) -> str:
    """
    读取 .doc 文件，采用轻量级双层兜底策略。

    解析链（按优先级）：antiword → python-ole
    所有解析器失败后抛出 DocParseError，绝不使用 GBK 二进制兜底。

    Raises:
        DocParseError: 所有解析器均失败时抛出
    """
    attempted = []

    def _try_decode(raw: bytes) -> str:
        """
        尝试多种编码解析 antiword 的原始输出。
        优先尝试 UTF-8（Git Bash 环境），再尝试 GBK/GB18030。
        返回解码后的字符串；若全部失败则返回空。
        """
        encodings = ['utf-8', 'gbk', 'gb18030']
        for enc in encodings:
            try:
                text = raw.decode(enc, errors='ignore')
                # 校验：必须有足够的中文字符（至少 30 个）
                chinese_count = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
                if chinese_count >= 30:
                    logger.debug(f"[_try_decode] succeeded with {enc}, chinese={chinese_count}")
                    return text
            except Exception:
                pass
        return ""

    # ── 优先级1：antiword ────────────────────────────────────────
    try:
        result = subprocess.run(
            ['antiword', '-t', '-w', '0', str(file_path)],
            capture_output=True, timeout=30
        )
        text = _try_decode(result.stdout)
        if result.returncode == 0 and len(text.strip()) > 50:
            logger.debug(f"[antiword] succeeded for {file_path}")
            return text.strip()
        else:
            logger.debug(
                f"[antiword] failed for {file_path}, "
                f"returncode={result.returncode}, stdout_len={len(result.stdout)}"
            )
            attempted.append('antiword')
    except FileNotFoundError:
        logger.debug(f"[antiword] not found for {file_path}")
        attempted.append('antiword (not available)')
    except subprocess.TimeoutExpired:
        logger.debug(f"[antiword] timeout for {file_path}")
        attempted.append('antiword (timeout)')

    # ── 优先级2：python-ole 兜底 ────────────────────────────────
    try:
        text = _extract_text_from_ole(str(file_path))
        if len(text.strip()) > 50:
            logger.debug(f"[python-ole] succeeded for {file_path}")
            return text.strip()
        attempted.append('python-ole (output too short)')
    except Exception as e:
        logger.debug(f"[python-ole] failed for {file_path}: {e}")
        attempted.append('python-ole (failed)')

    # ── 全部失败：抛出明确异常 ──────────────────────────────────
    logger.error(
        f"All .doc parsers failed for {file_path}",
        extra={'file': str(file_path), 'attempted_parsers': attempted}
    )
    raise DocParseError(str(file_path), attempted)

def parse_file(file_path: Path) -> str:
    """根据扩展名解析文件，调用对应解析器"""

    suffix = file_path.suffix.lower()
    if suffix == '.docx':
        content = read_docx(file_path)
    elif suffix == '.txt':
        content = read_txt(file_path)
    elif suffix == '.doc':
        content = read_doc(file_path)
    elif suffix == '.pdf':
        content = read_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    # 统一长度校验
    if len(content.strip()) < 50:
        raise DocumentParseError(f"Extracted text too short (<50 chars) for {file_path}")

    logger.info(f"Successfully parsed {file_path}, length={len(content)}")
    return content
